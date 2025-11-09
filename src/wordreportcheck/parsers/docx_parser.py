from pathlib import Path
from typing import List, Optional, Dict, Tuple
import re
from docx import Document

from ..schemas import ReportItem, ReportDocument


def _strip(text: Optional[str]) -> str:
    return (text or "").strip()


LABEL_MAP: Dict[str, str] = {
    "学院": "学院信息",
    "学院信息": "学院信息",
    "专业": "专业信息",
    "专业信息": "专业信息",
    "时间": "时间",
    "姓名": "姓名",
    "学号": "学号",
    "班级": "班级",
    "指导老师": "指导老师",
    "课程名称": "课程名称",
    "周次": "周次",
    "实验名称": "实验名称",
    "实验环境": "实验环境",
    "实验内容": "实验内容",
    "实验分析与体会": "实验分析与体会",
    "实验日期": "实验日期",
    "备注": "备注",
    "成绩": "成绩",
    "签名": "签名",
    "日期": "日期",
}


def _normalize_label(text: str) -> str:
    t = _strip(text)
    # 去掉末尾的冒号/全角冒号及空白
    t = re.sub(r"[：:]\s*$", "", t)
    # 替换常见空格/制表符
    t = re.sub(r"\s+", "", t)
    return LABEL_MAP.get(t, t)


def _extract_row_label_value(row) -> Optional[Dict[str, str]]:
    cells = row.cells
    if not cells:
        return None
    if len(cells) >= 2:
        label = _normalize_label(cells[0].text)
        value = _strip(cells[1].text)
        return {"label": label, "value": value}
    # 单列情况：尝试按冒号分割
    raw = _strip(cells[0].text)
    parts = re.split(r"[：:]", raw, maxsplit=1)
    if len(parts) == 2:
        label = _normalize_label(parts[0])
        value = _strip(parts[1])
        return {"label": label, "value": value}
    return None


def _parse_content_items(content_text: str, expected_count: int) -> List[ReportItem]:
    """将“实验内容”按题目段落切分为指定数量，并在每题中进一步提取：
    - 题目/题目要求
    - 实验方法和步骤（或 方法和步骤）
    - 代码

    规则与启发：
    - 题目段边界：按“题目N：”识别，N 可为阿拉伯数字或中文数字；标题到下一题目之间视为该题内容。
    - 段内标签：支持“题目要求/题目”、“实验方法和步骤/方法和步骤”、“代码”，允许前缀“题目N”。
    - 忽略“运行结果/结果展示”等不计入上述三类字段。
    - 数量保证：如少于 expected_count，则补空白题目；多于则截断。
    """
    text = content_text.replace("\r\n", "\n").replace("\r", "\n")

    # 题目段边界识别：允许中文数字与阿拉伯数字
    numeral = r"[一二三四五六七八九十百千零〇\d]+"
    seg_pat = re.compile(
        rf"(^|\n)(?:（[^）]*）)?\s*题目\s*{numeral}\s*[：:]\s*(.*?)(?=\n(?:（[^）]*）)?\s*题目\s*{numeral}\s*[：:]|\Z)",
        re.S,
    )

    segments: List[str] = []
    for m in seg_pat.finditer(text):
        segments.append(_strip(m.group(2)))

    # 若未识别到题目结构，退化为整段作为一个“未知题目”项（避免空输出）
    if not segments and content_text.strip():
        segments = [_strip(content_text)]

    def _extract_fields(seg: str) -> Tuple[Optional[str], str, Optional[str], Optional[str]]:
        """从单个题目段内提取 (title, question, methods, code)。
        - 支持标签前可有“题目N”前缀，如“题目1代码：”。
        - 若缺少明确标签：首行作为题目名称或题目要求；其余作为方法；代码为空。
        """
        # 查找标签位置
        label_pat = re.compile(
            rf"(?:^|\n)\s*(?:题目\s*{numeral}\s*)?(题目要求|题目|实验方法和步骤|方法和步骤|代码|运行结果)\s*[：:]\s*",
            re.I,
        )
        matches = list(label_pat.finditer(seg))
        title: Optional[str] = None
        question = ""
        methods: Optional[str] = None
        code: Optional[str] = None

        if matches:
            # 处理标签段：遍历标签对之间的内容
            # 先取第一个标签前的文本作为可能的标题
            first_start = matches[0].start()
            preamble = _strip(seg[:first_start])
            if preamble:
                # 取第一行作为题目名称
                title = _strip(preamble.split("\n", 1)[0])

            for i, m in enumerate(matches):
                label = m.group(1)
                start = m.end()
                end = matches[i + 1].start() if i + 1 < len(matches) else len(seg)
                content = _strip(seg[start:end])
                if not content:
                    continue
                if label in ("题目要求", "题目"):
                    question = content
                elif label in ("实验方法和步骤", "方法和步骤"):
                    methods = content
                elif label == "代码":
                    code = content
                elif label in ("运行结果",):
                    # 忽略运行结果
                    pass
        else:
            # 无标签：首行作为题目名称或要求，其余作为方法
            if "\n" in seg:
                first, rest = seg.split("\n", 1)
                title = _strip(first)
                question = title
                methods = _strip(rest)
            else:
                title = _strip(seg)
                question = title

        return title, question, methods, code

    items: List[ReportItem] = []
    for idx, seg in enumerate(segments):
        title, q, methods, code = _extract_fields(seg)
        parts = []
        if methods:
            parts.append(f"实验方法和步骤：{methods}")
        if code:
            parts.append(f"代码：{code}")
        answer = "\n".join(parts)
        # 题目编号按顺序生成；若 title 存在，保留在可选字段中
        try:
            items.append(ReportItem(id=f"Q{idx + 1}", question=q or "", answer=answer or "", methods=methods, code=code, title=title))
        except TypeError:
            items.append(ReportItem(id=f"Q{idx + 1}", question=q or "", answer=answer or ""))

    # 数量保证：少则补空，多则截断
    if expected_count is not None and expected_count > 0:
        if len(items) > expected_count:
            items = items[:expected_count]
        elif len(items) < expected_count:
            deficit = expected_count - len(items)
            base = len(items)
            for i in range(deficit):
                n = base + i + 1
                items.append(ReportItem(id=f"Q{n}", question="", answer=""))

    return items


def _get_grid_span(cell) -> int:
    # 兼容 python-docx 的 grid_span 属性，不存在时按 1 处理
    return getattr(cell, "grid_span", 1)


def _parse_by_template(doc: Document) -> Optional[ReportDocument]:
    """按照用户提供的统一模板（行号 + grid_span 顺序）解析 18 项字段。

    该模板特点：
    - 顶部若干行是信息汇总行，单元格可能合并，需用 grid_span 累积列索引
    - “实验内容”从包含该关键词的行开始，直到出现“实验分析与体会”为止
    - “实验分析与体会”之后的若干行依次为：实验日期、备注、成绩、（可能有签名行）、日期
    """
    if not doc.tables:
        return None

    table = doc.tables[0]
    rows = table.rows
    if not rows:
        return None

    report = ReportDocument(content_items=[])

    def set_field(name: str, value: str):
        try:
            setattr(report, name, _strip(value))
        except Exception:
            pass

    content_started = False
    analys_row_idx: Optional[int] = None
    content_buffer: List[str] = []

    for row_idx in range(len(rows)):
        row = rows[row_idx]
        # 行文本聚合（用于关键词检测）
        row_text = " ".join([_strip(c.text) for c in row.cells])

        # 先处理顶部信息汇总区（参考提供脚本的行号与索引）
        grid_span_index = 0
        col_idx = 0
        while col_idx < len(row.cells):
            cell = row.cells[col_idx]
            text = _strip(cell.text)
            span = _get_grid_span(cell)

            if row_idx == 0:
                if grid_span_index == 0:
                    set_field("学院信息", text)
                elif grid_span_index == 1:
                    set_field("专业信息", text)
                elif grid_span_index == 2:
                    set_field("时间", text)
            elif row_idx == 1:
                if grid_span_index == 1:
                    set_field("姓名", text)
                elif grid_span_index == 3:
                    set_field("学号", text)
            elif row_idx == 2:
                if grid_span_index == 1:
                    set_field("班级", text)
                elif grid_span_index == 3:
                    set_field("指导老师", text)
            elif row_idx == 3:
                if grid_span_index == 1:
                    set_field("课程名称", text)
                elif grid_span_index == 3:
                    set_field("周次", text)
            elif row_idx == 4:
                if grid_span_index == 1:
                    set_field("实验名称", text)

            # “实验分析与体会”之后的区块（以 analys_row_idx 为基准）
            if analys_row_idx is not None:
                if row_idx == analys_row_idx + 1:
                    set_field("实验分析与体会", text)
                elif row_idx == analys_row_idx + 2:
                    set_field("实验日期", text)
                elif row_idx == analys_row_idx + 3 and grid_span_index == 1:
                    set_field("备注", text)
                elif row_idx == analys_row_idx + 4 and grid_span_index == 1:
                    set_field("成绩", text)
                elif row_idx == analys_row_idx + 6 and grid_span_index == 1:
                    set_field("日期", text)

            col_idx += span
            grid_span_index += 1

        # 识别“实验内容”起止范围
        if not content_started and ("实验内容" in row_text):
            content_started = True
            continue

        if content_started and analys_row_idx is None:
            # 如果本行包含“实验分析与体会”，标记结束，并记录分析起始行索引
            if "实验分析与体会" in row_text:
                analys_row_idx = row_idx
            else:
                # 在内容区，累积文本（整行作为一个段）
                if row_text:
                    content_buffer.append(row_text)

    # 仅保存“实验内容”原文，不在解析阶段做分割
    if content_buffer:
        full_content = "\n\n".join(content_buffer)
        report.实验内容原文 = _strip(full_content)

    # 若至少有一个关键字段或内容则认为解析成功
    has_any = any([
        report.学院信息, report.专业信息, report.时间, report.姓名, report.学号,
        report.课程名称, report.实验名称, report.实验分析与体会, report.实验日期,
        (report.实验内容原文 and report.实验内容原文.strip()),
    ])
    return report if has_any else None


def parse_docx_to_report(docx_path: Path) -> ReportDocument:
    doc = Document(str(docx_path))
    # 优先按统一模板严格解析
    tmpl_report = _parse_by_template(doc)
    if tmpl_report:
        return tmpl_report

    # 模板不匹配时使用通用标签映射解析
    report = ReportDocument(content_items=[])
    content_accumulator: List[str] = []

    for table in doc.tables:
        for row in table.rows:
            pair = _extract_row_label_value(row)
            if not pair:
                continue
            label = pair["label"]
            value = pair["value"]
            if label == "实验内容":
                if value:
                    content_accumulator.append(value)
                continue
            # 仅当是我们识别的18个单元之一时赋值
            if label in LABEL_MAP.values():
                try:
                    setattr(report, label, value)
                except Exception:
                    pass

    # 仅保存“实验内容”原文，不在解析阶段做分割
    full_content = "\n\n".join(content_accumulator).strip()
    report.实验内容原文 = full_content if full_content else None

    # 题目分割不在解析阶段进行；由 CLI 在提供了分割数量时调用 _parse_content_items 完成

    return report